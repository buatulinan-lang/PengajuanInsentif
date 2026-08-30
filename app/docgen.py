"""Pembuat dokumen Word pengajuan insentif.

Dokumen disusun dari nol agar isinya selalu sesuai data pengajuan: kop cabang,
badan surat, tabel perhitungan sesuai jenis insentif, dan blok tanda tangan
dengan QR verifikasi pada kolom masing-masing pejabat.
"""
import io
import os
import qrcode
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.paths import STATIC_DIR

BULAN = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
         "Agustus", "September", "Oktober", "November", "Desember"]

NAVY = RGBColor(0x14, 0x35, 0x5B)
ABU = RGBColor(0x60, 0x6A, 0x76)
KREM = "EFE4D2"
KREM_MUDA = "F7F1E6"


def rupiah(n):
    return "Rp " + "{:,.0f}".format(float(n or 0)).replace(",", ".")


def angka(n):
    return "{:,.0f}".format(float(n or 0)).replace(",", ".")


def make_qr(payload, px=10):
    qr = qrcode.QRCode(box_size=px, border=1,
                       error_correction=qrcode.constants.ERROR_CORRECT_M)
    qr.add_data(payload)
    qr.make(fit=True)
    buf = io.BytesIO()
    qr.make_image(fill_color="black", back_color="white").save(buf, format="PNG")
    buf.seek(0)
    return buf


# ------------------------------------------------------------------ bantuan
def _arsir(sel, warna):
    tcPr = sel._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), warna)
    tcPr.append(sh)


def _garis_bawah(p, tebal=8, warna="14355B"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(tebal))
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), warna)
    pbdr.append(b)
    pPr.append(pbdr)


def _teks(sel, isi, ukuran=8.5, tebal=False, rata=None, warna=None):
    sel.text = ""
    p = sel.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if rata:
        p.alignment = rata
    r = p.add_run(str(isi))
    r.font.size = Pt(ukuran)
    r.bold = tebal
    if warna:
        r.font.color.rgb = warna
    return p


def _paragraf(doc, isi="", ukuran=10.5, tebal=False, rata=None, italic=False,
              spasi_setelah=6, warna=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spasi_setelah)
    p.paragraph_format.space_before = Pt(0)
    if rata:
        p.alignment = rata
    if isi:
        r = p.add_run(isi)
        r.font.size = Pt(ukuran)
        r.bold = tebal
        r.italic = italic
        if warna:
            r.font.color.rgb = warna
    return p


def _jangan_pecah(t, sampai=None):
    """Cegah baris tabel terbelah antar halaman; jaga blok tetap menyatu."""
    baris = t.rows if sampai is None else t.rows[:sampai]
    for r in t.rows:
        trPr = r._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit")
        trPr.append(cs)
    for r in baris:
        for sel in r.cells:
            for par in sel.paragraphs:
                par.paragraph_format.keep_with_next = True


def _lebar_tetap(t, lebar=None):
    """Kunci lebar kolom agar tidak diatur ulang oleh Word.

    Lebar sel saja tidak cukup: Word dan LibreOffice memakai w:tblGrid pada
    tata letak tetap, sedangkan python-docx tidak ikut memperbaruinya.
    """
    tblPr = t._tbl.tblPr
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    if lebar:
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for kol in list(grid):
                grid.remove(kol)
            for w in lebar:
                kol = OxmlElement("w:gridCol")
                kol.set(qn("w:w"), str(int(round(w * 567))))   # cm -> twips
                grid.append(kol)


def _tabel(doc, judul, lebar, ukuran=8.5):
    t = doc.add_table(rows=1, cols=len(judul))
    t.style = "Table Grid"
    t.autofit = False
    _lebar_tetap(t, lebar)
    for i, (j, w) in enumerate(zip(judul, lebar)):
        sel = t.rows[0].cells[i]
        sel.width = Cm(w)
        _arsir(sel, KREM)
        _teks(sel, j, ukuran, True, WD_ALIGN_PARAGRAPH.CENTER, NAVY)
    return t


def _baris_lebar(t, label, nilai, lebar, ukuran=8.5, tebal=False, arsir=None):
    """Baris dengan label melebar ke seluruh kolom kecuali kolom nilai terakhir."""
    r = t.add_row()
    for i, w in enumerate(lebar):
        r.cells[i].width = Cm(w)
    kiri = r.cells[0].merge(r.cells[len(lebar) - 2])
    kiri.width = Cm(sum(lebar[:-1]))
    if arsir:
        _arsir(kiri, arsir)
        _arsir(r.cells[-1], arsir)
    _teks(kiri, label, ukuran, tebal)
    _teks(r.cells[-1], nilai, ukuran, tebal, WD_ALIGN_PARAGRAPH.RIGHT)
    return r


def _baris(t, nilai, lebar, ukuran=8.5, tebal=False, arsir=None, kanan=()):
    r = t.add_row()
    for i, (v, w) in enumerate(zip(nilai, lebar)):
        sel = r.cells[i]
        sel.width = Cm(w)
        sel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if arsir:
            _arsir(sel, arsir)
        rata = WD_ALIGN_PARAGRAPH.RIGHT if i in kanan else None
        _teks(sel, v, ukuran, tebal, rata)
    return r


# ------------------------------------------------------------------ kop surat
def _kop(doc, cabang):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(1.6), Cm(1.8)
    sec.left_margin, sec.right_margin = Cm(2.2), Cm(2.2)
    sec.header_distance = Cm(1.0)

    hdr = sec.header
    hdr.is_linked_to_previous = False
    for p in list(hdr.paragraphs):
        p.text = ""

    t = hdr.add_table(rows=1, cols=2, width=Cm(16.6))
    t.autofit = False
    kiri, kanan = t.rows[0].cells
    kiri.width, kanan.width = Cm(3.2), Cm(13.4)

    logo = os.path.join(STATIC_DIR, "logo-mflash.png")
    kiri.text = ""
    if os.path.exists(logo):
        kiri.paragraphs[0].add_run().add_picture(logo, height=Cm(1.35))

    kanan.text = ""
    p = kanan.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(cabang.display_name or f"MFlash – {cabang.name}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    p2 = kanan.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(cabang.address or "")
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = ABU

    garis = hdr.add_paragraph()
    garis.paragraph_format.space_before = Pt(2)
    garis.paragraph_format.space_after = Pt(0)
    _garis_bawah(garis, 12, "F58F22")


def _kepala_surat(doc, sub, judul, kepada, tembusan):
    _paragraf(doc, judul.upper(), 13, True, WD_ALIGN_PARAGRAPH.CENTER,
              spasi_setelah=2, warna=NAVY)
    _paragraf(doc, f"Nomor: {sub.code}", 8.5, False, WD_ALIGN_PARAGRAPH.CENTER,
              spasi_setelah=12, warna=ABU)

    t = doc.add_table(rows=0, cols=2)
    t.autofit = False
    for label, isi in (("Lampiran", "1 (satu) berkas"),
                       ("Dari", sub.submitter.position or "Store Leader"),
                       ("Kepada", kepada),
                       ("Tembusan", tembusan)):
        r = t.add_row()
        r.cells[0].width, r.cells[1].width = Cm(2.6), Cm(14.0)
        _teks(r.cells[0], label, 10)
        _teks(r.cells[1], ": " + isi, 10)
    _paragraf(doc, spasi_setelah=6)


def _pembuka(doc, sub, kalimat):
    _paragraf(doc, "Bismillaahirrahmaanirrahiim", 10.5, False, italic=True,
              spasi_setelah=2)
    _paragraf(doc, "Assalaamu'alaikum Warahmatullaahi Wabarakaatuh", 10.5,
              spasi_setelah=8)
    _paragraf(doc, kalimat, 10.5, spasi_setelah=8)
    _paragraf(doc, "Adapun rincian perhitungannya sebagai berikut:", 10.5,
              spasi_setelah=8)


def _penutup_ttd(doc, sub, approvals, base_url):
    _paragraf(doc, spasi_setelah=4)
    _paragraf(doc, "Demikian yang dapat kami sampaikan. Atas perhatian dan "
                   "kerja samanya kami ucapkan terima kasih.", 10.5, spasi_setelah=4)
    _paragraf(doc, "Wassalaamu'alaikum Warahmatullaahi Wabarakaatuh", 10.5,
              spasi_setelah=10)

    kota = sub.branch.city if sub.branch and sub.branch.city else "Jakarta"
    tgl = sub.created_at
    ptgl = _paragraf(doc, f"{kota}, {tgl.day} {BULAN[tgl.month]} {tgl.year}", 10.5,
                     rata=WD_ALIGN_PARAGRAPH.RIGHT, spasi_setelah=10)
    ptgl.paragraph_format.keep_with_next = True

    # peta persetujuan menurut peran
    oleh = {}
    for a in approvals:
        if a.action in ("submit", "approve"):
            oleh[a.role] = a

    if sub.type == "profit_arm":
        kolom = [("Diajukan oleh,", "arm"),
                 ("Disetujui oleh,", "ceo"),
                 ("Disetujui oleh,", "finance")]
    else:
        kolom = [("Diajukan oleh,", "store_leader"),
                 ("Diperiksa oleh,", "arm"),
                 ("Disetujui oleh,", "ceo"),
                 ("Diverifikasi oleh,", "finance")]
    t = doc.add_table(rows=4, cols=len(kolom))
    t.autofit = False
    lebar = Cm(16.6 / len(kolom))

    for i, (label, peran) in enumerate(kolom):
        a = oleh.get(peran)
        for baris in range(4):
            t.rows[baris].cells[i].width = lebar

        _teks(t.rows[0].cells[i], label, 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)

        sel = t.rows[1].cells[i]
        sel.text = ""
        p = sel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        if a:
            p.add_run().add_picture(make_qr(f"{base_url}/verify/{a.qr_token}"),
                                    width=Cm(2.1))
        else:
            r = p.add_run("\n\n(menunggu)\n")
            r.font.size = Pt(8)
            r.font.color.rgb = ABU

        nama = a.user.full_name if a else "—"
        pn = _teks(t.rows[2].cells[i], nama, 9.5, True, WD_ALIGN_PARAGRAPH.CENTER)
        pn.runs[0].underline = True

        jab = (a.user.position if a and a.user.position else
               {"store_leader": "Store Leader", "arm": "Area Operation Manager",
                "ceo": "CEO MFlash", "finance": "Finance Manager"}[peran])
        _teks(t.rows[3].cells[i], jab, 8.5, False, WD_ALIGN_PARAGRAPH.CENTER, ABU)

    _jangan_pecah(t, sampai=3)   # empat baris tanda tangan tidak terpisah halaman

    _paragraf(doc, spasi_setelah=4)
    p = _paragraf(doc,
                  "Dokumen ini disetujui secara elektronik. Pindai QR pada kolom "
                  "tanda tangan untuk memverifikasi nama, jabatan, dan waktu "
                  "persetujuan.", 7.5, rata=WD_ALIGN_PARAGRAPH.CENTER, warna=ABU)
    _garis_bawah(p, 4, "E5D9C6")


# ------------------------------------------------------------------ isi tabel
def _tabel_profit(doc, hasil):
    judul = ["Cabang", "Omzet", "Laba Kotor", "Biaya\nOperasional",
             "Laba Bersih", "% thd\nOmzet", "Pengali", "Insentif"]
    lebar = [2.5, 2.5, 2.3, 2.2, 2.4, 1.5, 1.3, 2.0]
    t = _tabel(doc, judul, lebar)
    kanan = {1, 2, 3, 4, 5, 6, 7}

    for bl in hasil.get("blok", []):
        sisi = [("asal", bl["asal"])]
        if bl["tipe"] == "rotasi":
            sisi.append(("tujuan", bl["tujuan"]))
        for peran, d in sisi:
            h = d["hasil"]
            label = d["cabang"]
            if bl["tipe"] == "rotasi":
                label += "\n(cabang " + ("asal" if peran == "asal" else "tujuan") + ")"
            _baris(t, [label, angka(h["omzet"]), angka(h["laba_kotor_sesudah_fee"]),
                       angka(h["biaya_operasional"]), angka(h["laba_bersih"]),
                       f"{h['gp_pct']}%", f"{d['pengali']:.0f}%",
                       angka(d["insentif"])], lebar, kanan=kanan)
        if bl["tipe"] == "rotasi":
            _baris_lebar(t, f"Rotasi efektif {bl['mutasi']} — pengajuan jatuh pada "
                            f"{bl['offset_label']}. {bl['kasus']}.", "", lebar,
                         7.5, arsir=KREM_MUDA)

    _baris_lebar(t, "Subtotal insentif seluruh cabang",
                 angka(hasil.get("subtotal_blok", hasil.get("total", 0))),
                 lebar, 8.5, True, KREM_MUDA)
    for p in hasil.get("rincian_potongan", []):
        _baris_lebar(t, f"Pengurang {p['pct']:.0f}% — {p['nama']}",
                     "− " + angka(p["nilai"]), lebar, 8, arsir=KREM_MUDA)
    _baris_lebar(t, "TOTAL INSENTIF YANG DIDAPAT", angka(hasil.get("total", 0)),
                 lebar, 9.5, True, KREM)
    _catatan(doc, "Laba Bersih dihitung dari Total Laba Bersih Setelah Prepaid "
                  "dikurangi Laba Ditahan. Besaran insentif mengikuti matriks "
                  "TAP CEO yang berlaku.")

    _tabel_goal(doc, hasil)


def _tabel_sales(doc, hasil):
    judul = ["Nama Sales", "Status", "Omzet", "GP Aksesoris",
             "Unit HP", "Unit Laptop", "Insentif"]
    lebar = [4.8, 1.8, 2.5, 2.3, 1.5, 1.6, 2.1]
    t = _tabel(doc, judul, lebar, 8)
    kanan = {2, 3, 4, 5, 6}
    semua = hasil.get("baris", [])
    # Sales tanpa insentif bulan ini tidak dicetak supaya daftarnya ringkas.
    aktif = [b for b in semua if b.get("total")]
    for b in aktif:
        _baris(t, [b["nama"], b["status"], angka(b["omset_total"]),
                   angka(b["gp_aksesoris"]), b["n_handphone"], b["n_laptop"],
                   angka(b["total"])], lebar, 8, kanan=kanan)
    _baris_lebar(t, "Subtotal Insentif Sales", angka(hasil.get("subtotal_sales", 0)),
                 lebar, 8.5, True, KREM_MUDA)
    _baris_lebar(t, f"Insentif Team {hasil.get('pct_insentif_team', 2)}% dari bagi "
                    f"hasil service MFlash pelanggan Member Reguler "
                    f"({angka(hasil.get('omset_service_member', 0))})",
                 angka(hasil.get("insentif_team", 0)), lebar, 8.5, False, KREM_MUDA)
    _baris_lebar(t, "TOTAL INSENTIF YANG DIDAPAT", angka(hasil.get("total", 0)),
                 lebar, 9.5, True, KREM)
    nihil = len(semua) - len(aktif)
    _catatan(doc, "Omzet dan gross profit diatribusikan melalui Nama Default "
                  "Penjual pada data pelanggan. Jumlah unit dihitung per faktur "
                  "menurut kategori penjualan."
                  + (f" {nihil} sales tanpa insentif bulan ini tidak ditampilkan."
                     if nihil else ""))


def _tabel_purchasing(doc, hasil):
    lebar = [10.6, 6.0]
    t = _tabel(doc, ["Uraian", "Nilai"], lebar, 9)
    baris = [
        ("Pembelian ke supplier TERTARGET",
         f"{angka(hasil['nilai_tertarget'])}   ({hasil['pct_tertarget']}%)"),
        ("Pembelian ke supplier NON TERTARGET",
         angka(hasil["nilai_non_tertarget"])),
        ("Total Pembelian", angka(hasil["total_pembelian"])),
        ("Tiering pencapaian goal", f"{hasil['tier_pct']}%"),
        ("Nominal insentif", angka(hasil["nominal_insentif"])),
        (f"Pencapaian terhadap Goal Performate {hasil['goal_pct']}%",
         f"{hasil['pencapaian_goal']}%"),
        ("Pengali berdasarkan supplier tertarget", f"{hasil['pengali_pct']}%"),
    ]
    for a, b in baris:
        _baris(t, [a, b], lebar, 9, kanan={1})
    _baris(t, ["INSENTIF FINAL", angka(hasil["total"])], lebar, 10, True, KREM,
           kanan={1})



def _tabel_arm(doc, hasil):
    judul = ["Cabang", "Omzet", "Laba Bersih Setelah Prepaid",
             "Laba Ditahan", "Laba Bersih"]
    lebar = [4.2, 3.3, 3.6, 2.8, 3.3]
    t = _tabel(doc, judul, lebar)
    kanan = {1, 2, 3, 4}
    for b in hasil.get("baris", []):
        nama = b["cabang"] + ("" if b["ikut"] else "  (tidak dihitung)")
        if b.get("galat"):
            _baris(t, [nama, "—", "—", "—", "—"], lebar, 8, kanan=kanan)
            continue
        _baris(t, [nama, angka(b["omzet"]), angka(b["laba_setelah_prepaid"]),
                   angka(b["laba_ditahan"]), angka(b["laba_bersih"])],
               lebar, 8.5, arsir=None if b["ikut"] else KREM_MUDA, kanan=kanan)
    _baris_lebar(t, "TOTAL LABA BERSIH SELURUH CABANG",
                 angka(hasil.get("total_laba_bersih", 0)), lebar, 9, True, KREM)
    _catatan(doc, f"Laba Bersih tiap cabang = Total Laba Bersih Setelah Prepaid "
                  f"dikurangi Laba Ditahan {hasil.get('laba_ditahan_pct', 7)}%. "
                  f"Cabang bertanda \u201ctidak dihitung\u201d tidak menambah total.")

    _paragraf(doc, spasi_setelah=8)
    _paragraf(doc, "Perhitungan Insentif", 10, True, spasi_setelah=4, warna=NAVY)
    lg = [10.6, 6.0]
    tg = _tabel(doc, ["Uraian", "Nilai"], lg, 9)
    _baris(tg, ["Total laba bersih seluruh cabang",
                angka(hasil["total_laba_bersih"])], lg, 9, kanan={1})
    _baris(tg, [f"Baris matriks (dibulatkan ke bawah)",
                angka(hasil["baris_matriks"] or 0)], lg, 9, kanan={1})
    _baris(tg, ["Nominal insentif menurut matriks",
                angka(hasil["nominal_matriks"])], lg, 9, kanan={1})
    pp = hasil.get("pencapaian_prioritas")
    _baris(tg, [f"Pencapaian Target Prioritas"
                f"{'' if pp is None else f' ({pp}%)'} — {hasil['label_pengali']}",
                f"{hasil['pengali_pct']}%"], lg, 9, kanan={1})
    _baris(tg, ["Insentif setelah pengali", angka(hasil["setelah_pengali"])],
           lg, 9, kanan={1})
    for p in hasil.get("rincian_potongan", []):
        _baris(tg, [f"Pengurang {p['pct']:.0f}% — {p['nama']}",
                    "− " + angka(p["nilai"])], lg, 8.5, arsir=KREM_MUDA, kanan={1})
    _baris(tg, ["INSENTIF YANG DIDAPAT", angka(hasil["total"])], lg, 10, True,
           KREM, kanan={1})


def _tabel_goal(doc, hasil):
    diisi = [g for g in hasil.get("goal", []) if g.get("diisi")]
    if not diisi:
        return
    _paragraf(doc, spasi_setelah=6)
    _paragraf(doc, "Pencapaian Goal", 10, True, spasi_setelah=4, warna=NAVY)
    lg = [8.6, 4.0, 4.0]
    tg = _tabel(doc, ["Goal", "Pencapaian", "Status"], lg, 8.5)
    for g in diisi:
        _baris(tg, [g["nama"] + (" (utama)" if g["utama"] else ""),
                    f"{g['pencapaian']}%",
                    "Tercapai" if g["tercapai"] else
                    f"Di bawah {hasil.get('ambang_goal_pct', 98)}%"],
               lg, 8.5, kanan={1})


def _catatan(doc, isi):
    _paragraf(doc, isi, 7.5, spasi_setelah=2, warna=ABU, italic=True)


# ------------------------------------------------------------------ utama
JUDUL = {
    "profit_sl": "Pengajuan Insentif Profit Store Leader",
    "sales_team": "Pengajuan Insentif Sales dan Team",
    "purchasing": "Pengajuan Insentif Purchasing",
    "profit_arm": "Pengajuan Insentif Profit ARM",
}


def buat_dokumen(sub, hasil, approvals, base_url, out_path):
    doc = Document()
    gaya = doc.styles["Normal"]
    gaya.font.name = "Calibri"
    gaya.font.size = Pt(10.5)
    gaya.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    _kop(doc, sub.branch)

    periode = f"{BULAN[sub.period_month]} {sub.period_year}"
    total = rupiah(hasil.get("total", sub.total_amount))

    if sub.type == "sales_team":
        kepada = "Finance Manager"
        tembusan = "CEO Madinah Group Indonesia, Corporate Strategy, Store Area Manager"
        kalimat = (f"Sehubungan dengan telah berakhirnya periode bulan {periode}, "
                   f"maka dengan ini kami mengajukan Komisi Penjualan Member Sales "
                   f"dan Team MFlash {sub.branch.name} senilai {total} berdasarkan "
                   f"data penjualan periode tersebut, sesuai ketetapan yang berlaku.")
    elif sub.type == "profit_arm":
        kepada = "Finance Manager"
        tembusan = "CEO Madinah Group Indonesia, CEO MFlash, HR Manager"
        kalimat = (f"Sehubungan dengan telah berakhirnya periode bulan {periode}, "
                   f"maka dengan ini saya mengajukan Insentif Profit Area Manager "
                   f"MFlash senilai {total} berdasarkan laporan keuangan seluruh "
                   f"cabang pada periode tersebut, sesuai ketetapan yang berlaku.")
    elif sub.type == "purchasing":
        kepada = "Finance Manager"
        tembusan = "CEO Madinah Group Indonesia, Area Operation Manager"
        kalimat = (f"Sehubungan dengan telah berakhirnya periode bulan {periode}, "
                   f"maka dengan ini kami mengajukan Insentif Purchasing MFlash "
                   f"{sub.branch.name} senilai {total} dengan data sebagai berikut.")
    else:
        kepada = "Finance Manager"
        tembusan = "CEO Madinah Group Indonesia, Corporate Strategy, HR Manager"
        kalimat = (f"Sehubungan dengan telah berakhirnya periode bulan {periode}, "
                   f"maka dengan ini saya mengajukan Insentif Bulanan senilai "
                   f"{total} sesuai dengan TAP CEO yang berlaku.")

    _kepala_surat(doc, sub, JUDUL.get(sub.type, "Pengajuan Insentif"),
                  kepada, tembusan)
    _pembuka(doc, sub, kalimat)

    if not hasil:
        # Perhitungan gagal (mis. master data belum lengkap): cetak keterangannya
        # supaya dokumen tetap terbuka, bukan gagal diunduh.
        _paragraf(doc, "Rincian perhitungan belum tersedia.", 10.5, True,
                  spasi_setelah=4)
        _paragraf(doc, sub.note or "Data pengajuan belum lengkap. Silakan buat "
                                   "ulang pengajuan setelah data pendukungnya "
                                   "dilengkapi.", 10, spasi_setelah=10, warna=ABU)
    elif sub.type == "sales_team":
        _tabel_sales(doc, hasil)
    elif sub.type == "purchasing":
        _tabel_purchasing(doc, hasil)
    elif sub.type == "profit_arm":
        _tabel_arm(doc, hasil)
        _tabel_goal(doc, hasil)
    else:
        _tabel_profit(doc, hasil)

    _penutup_ttd(doc, sub, approvals, base_url)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
